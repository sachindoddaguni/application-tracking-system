"""
The flask application for our program
"""
# importing required python libraries
from flask import Flask, jsonify, request, send_file
from flask_mongoengine import MongoEngine
from flask_cors import CORS, cross_origin
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from bs4 import BeautifulSoup
from webdriver_manager.chrome import ChromeDriverManager
import pandas as pd
from datetime import datetime, timedelta
import yaml
import hashlib
import uuid
import os
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json

import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from apscheduler.schedulers.background import BackgroundScheduler
from datetime import datetime, timedelta
from collections import defaultdict
import sys

import json
import bcrypt
from flask_pymongo import PyMongo
from flask import Flask, request, jsonify
from flask_pymongo import PyMongo
from pymongo import MongoClient
import mongomock
from datetime import datetime, timedelta, timezone
from flask_jwt_extended import create_access_token,get_jwt,get_jwt_identity, \
                               unset_jwt_cookies, jwt_required, JWTManager
from datetime import datetime, timedelta
from functools import reduce
from bson import json_util 
from pymongo import MongoClient
from flasgger import Swagger




existing_endpoints = ["/applications", "/resume", "/dashboard", "/contacts", "/token", "/register"]




def create_app():
    """
    Creates a server hosted on localhost

    :return: Flask object
    """
    app = Flask(__name__)
    app.secret_key = 'secret'
    app.config["JWT_SECRET_KEY"] = "softwareEngineering"
    app.config["JWT_ACCESS_TOKEN_EXPIRES"] = timedelta(hours=1)
    jwt = JWTManager(app)
    # make flask support CORS
    CORS(app)
    app.config["CORS_HEADERS"] = "Content-Type"

    @app.route("/")
    @cross_origin()
    def health_check():
        return jsonify({"message": "Server up and running"}), 200
     

    @app.route('/token', methods=["POST"])
    def create_token():
        print('This is standard output')
        try:
            email = request.json.get("email", None)
            password = request.json.get("password", None)
            print(password)
            user = Users.objects(
                email=email, password=password
            ).first()
            if user is not None:
                access_token = create_access_token(identity=email)
                return jsonify({"message": "Login successful", "access_token":access_token})
            else:
                print("Invalid email or password")
                return jsonify({"message": "Invalid email or password"}),401
        except Exception as err:
            print( err)
            return jsonify({"error": "Internal server error"}), 500

        
    @app.route("/register", methods=["POST"])
    def register():
        email = request.json.get('email', None)
        password = request.json.get('password', None)
        first_name = request.json.get('firstName', None)
        last_name = request.json.get('lastName', None)
        print(email)
        new_document = {
        "email": email,
        "password": password,
        "first_name": first_name,
        "last_name": last_name,
        }
        query = {
            "email": email,
        }
        try:
            user = Users(
                firstName=first_name,
                lastName=last_name,
                email=email,
                password=password,
                applications=[]
            )
            user.save()
            return jsonify(user.to_json()), 200
        except Exception as e:
            response = jsonify({"msg": "register failed"})

        return response


        
    @app.route("/logout", methods=["POST"])
    def logout():
        """
        Logout the user and clear their session

        ---
        tags:
        - User Logout
        responses:
        200:
            description: Logout successful
        """
        response = jsonify({"msg": "logout successful"})
        unset_jwt_cookies(response)
        return response
    


    # get data from the CSV file for rendering root page
    @app.route("/applications", methods=["GET"])
    @jwt_required()
    def get_data():
        """
        Gets user's applications data from the database

        :return: JSON object with application data
        """
        print("In get data")
        try:
            user_email = get_jwt_identity()
            users = Users.objects()
            user = users.filter(email=user_email).first()
            applications = user["applications"]
            return jsonify(applications)
        except Exception as e:
            print(e)
            return jsonify({"error": "Internal server error"}), 500

    @app.route("/applications", methods=["POST"])
    @jwt_required()
    def add_application():
        """
        Add a new job application for the user

        :return: JSON object with status and message
        """
        
        try:
            current_user_email = get_jwt_identity()
            print("User: ------ ", current_user_email)
            users = Users.objects()
            user = users.filter(email=current_user_email).first()
            print("User: ------ ", user.email)
            current_application = {
                "id": get_new_application_id(user),
                "jobTitle": request.json.get('jobTitle', None),
                "companyName": request.json.get('companyName', None),
                "date": request.json.get('date', None),
                "jobLink": request.json.get('jobLink', None),
                "location": request.json.get('location', None),
                "stage": request.json.get("status", "1"),
            }
            print(current_application)
            applications = user["applications"] + [current_application]
            user["applications"] = applications
            user.save()
            return jsonify(current_application), 200
        except Exception as ex:
            return jsonify({"error": ex}), 500

    @app.route("/applications/<int:application_id>", methods=["PUT"])
    def update_application(application_id):
        """
        Updates the existing job application for the user

        :param application_id: Application id to be modified
        :return: JSON object with status and message
        """
        try:
            user_email = get_jwt_identity()
            try:
                request_data = json.loads(request.data)["application"]
            except:
                return jsonify({"error": "No fields found in input"}), 400

            users = Users.objects()
            user = users.filter(email=user_email).first()
            current_applications = user["applications"]

            if len(current_applications) == 0:
                return jsonify({"error": "No applications found"}), 400
            else:
                updated_applications = []
                app_to_update = None
                application_updated_flag = False
                for application in current_applications:
                    if application["id"] == application_id:
                        app_to_update = application
                        application_updated_flag = True
                        for key, value in request_data.items():
                            application[key] = value
                    updated_applications += [application]
                    try:
                    # Send an email notification
                        to_email = user.email  # Use the user's email address
                        subject = "Job Application Updated"
                        message = f"Hello {user.fullName},\n\nThe following job application has been updated:\nJob Title: {application['jobTitle']}\nCompany: {application['companyName']}\n\nBest regards,\nYour Application Tracker"
                        send_email(to_email, subject, message)
                    except:
                        return jsonify({"error": "EMAIL wasn't sent"}), 400
                if not application_updated_flag:
                    return jsonify({"error": "Application not found"}), 400
                
                user.update(applications=updated_applications)

            return jsonify(app_to_update), 200
        except:
            return jsonify({"error": "Internal server error"}), 500

    @app.route("/applications/<int:application_id>", methods=["DELETE"])
    def delete_application(application_id):
        """
        Deletes the given job application for the user

        :param application_id: Application id to be modified
        :return: JSON object with status and message
        """
        try:
            user_email = get_jwt_identity()
            users = Users.objects()
            user = users.filter(email=user_email).first()

            current_applications = user["applications"]

            application_deleted_flag = False
            updated_applications = []
            app_to_delete = None
            for application in current_applications:
                if application["id"] != application_id:
                    updated_applications += [application]
                else:
                    app_to_delete = application
                    application_deleted_flag = True
                    try:
                        # Send an email notification
                        to_email = user.email  # Use the user's email address
                        subject = "Job Application Deleted"
                        message = f"Hello {user.fullName},\n\nThe following job application has been deleted:\nJob Title: {application['jobTitle']}\nCompany: {application['companyName']}\n\nBest regards,\nYour Application Tracker"
                        send_email(to_email, subject, message)
                    except:
                            return jsonify({"error": "EMAIL wasn't sent"}), 400

            if not application_deleted_flag:
                return jsonify({"error": "Application not found"}), 400
            user.update(applications=updated_applications)
            return jsonify(app_to_delete), 200
        except:
            return jsonify({"error": "Internal server error"}), 500

    @app.route("/resume", methods=["POST"])
    def upload_resume():
        """
        Uploads resume file or updates an existing resume for the user

        :return: JSON object with status and message
        """
        try:
            user_email = get_jwt_identity()
            try:
                file = request.files["file"].read()
            except:
                return jsonify({"error": "No resume file found in the input"}), 400

            users = Users.objects()
            user = users.filter(email=user_email).first()
            if not user.resume.read():
                # There is no file
                user.resume.put(file)
                user.save()
                return jsonify({"message": "resume successfully uploaded"}), 200
            else:
                # There is a file, we are replacing it
                user.resume.replace(file)
                user.save()
                return jsonify({"message": "resume successfully replaced"}), 200
        except Exception as e:
            print(e)
            return jsonify({"error": "Internal server error"}), 500

    # @app.route("/resume", methods=["GET"])
    # def get_resume():
    #     """
    #     Retrieves the resume file for the user

    #     :return: response with file
    #     """
    #     try:
    #         userid = get_userid_from_header()
    #         try:
    #             user = Users.objects(id=userid).first()
    #             if len(user.resume.read()) == 0:
    #                 raise FileNotFoundError
    #             else:
    #                 user.resume.seek(0)
    #         except:
    #             return jsonify({"error": "resume could not be found"}), 400

    #         response = send_file(
    #             user.resume,
    #             mimetype="application/pdf",
    #             attachment_filename="resume.pdf",
    #             as_attachment=True,
    #         )
    #         response.headers["x-filename"] = "resume.pdf"
    #         response.headers["Access-Control-Expose-Headers"] = "x-filename"
    #         return response, 200
    #         return jsonify({"error": "Internal server error"}), 500

    @app.route("/dashboard", methods=["GET"])
    def get_dashboard_data():
        """
        Gets user's stats data from the database

        :return: JSON object with stats data
        """
        try:
            user_email = get_jwt_identity()
            users = Users.objects()         
            user = users.filter(email=user_email).first()
            applications = user["applications"]
            job_app_status, applications_created, interviews_completed = get_job_app_status(applications)
            six_months_job_count = get_last_six_months_job_counts(applications)
            last_four_apps = get_last_four_jobs(applications)
            contacts_saved = 10
            notes_taken = 15
            return jsonify({
                "six_months_jobs_count": six_months_job_count,
                "job_applications_status": job_app_status,
                "applications_created": applications_created,
                "interviews_completed": interviews_completed,
                "contacts_saved": contacts_saved,
                "notes_taken": notes_taken,
                "last_four_apps": last_four_apps
            }),200
        except Exception as e:
            print(e)
            return jsonify({"error": "Internal server error"}), 500
        

    @app.route("/users/contacts", methods=["GET"])
    @jwt_required()
    def get_contacts():
        """
        Retrieves contacts for the logged-in user.

        :return: JSON object with user's contacts
        """
        try:
            current_user = get_jwt_identity()
            user = Users.objects(email=current_user).first()
            if not user:
                return jsonify({"error": "User not found"}), 404
            contacts = user.contacts
            return jsonify({"contacts": contacts}), 200
        except Exception as e:
            return jsonify({"error": "Unable to fetch contacts"}), 500


    @app.route("/users/contacts", methods=["POST"])
    @jwt_required()
    def add_contact():
        """
        Adds a new contact for the user.

        :return: JSON object with status and message
        """
        try:
            current_user_id = get_jwt_identity()
            user = Users.objects(email=current_user_id).first()
            if not user:
                return jsonify({"error": "User not found"}), 404

            data = json.loads(request.data)
            new_contact = {
                "firstName": data["firstName"],
                "lastName": data["lastName"],
                "jobTitle": data.get("jobTitle", ""),
                "company": data.get("company", ""),
                "email": data.get("email", ""),
                "phone": data.get("phone", ""),
                "linkedin": data.get("linkedin", "")
            }
            user.contacts.append(new_contact)
            user.save()
            return jsonify({"message": "Contact added successfully", "contact": new_contact}), 200
        except Exception as e:
            print(e)
            return jsonify({"error": "Unable to add contact"}), 500
        
    return app


app = create_app()
# with open("application.yml") as f:
#     info = yaml.load(f, Loader=yaml.FullLoader)
#     username = info["username"]
#     password = info["password"]
#     app.config["MONGODB_SETTINGS"] = {
#         "db": "appTracker",
#         # "host": os.getenv("db_username"),
#         "host": "localhost",
#     }

app.config["MONGODB_SETTINGS"] = {
        "db": "appTracker",
        # "host": os.getenv("db_username"),
        "host": "localhost",
    }
db = MongoEngine()
db.init_app(app)


class Users(db.Document):
    """
    Users class. Holds full name, username, password, as well as applications and resumes
    """

    firstName = db.StringField()
    lastName = db.StringField()
    email    = db.StringField()
    password = db.StringField()
    applications = db.ListField()
    contacts = db.ListField()

    def to_json(self):
        """
        Returns the user details in JSON object

        :return: JSON object
        """
        return { "fullName": self.firstName, "lastName": self.lastName,  "email": self.email}


def get_new_user_id():
    """
    Returns the next value to be used for new user

    :return: key with new user_id
    """
    user_objects = Users.objects()
    if len(user_objects) == 0:
        return 1

    new_id = 0
    for a in user_objects:
        new_id = max(new_id, a["id"])

    return new_id + 1


def get_new_application_id(user):
    """
    Returns the next value to be used for new application

    :param: user_id: User id of the active user
    :return: key with new application_id
    """

    if len(user["applications"]) == 0:
        return 1

    new_id = 0
    for a in user["applications"]:
        new_id = max(new_id, a["id"])

    return new_id + 1

def send_email(to_email, subject, message):
    # Set up your email and password here, or use environment variables
    gmail_user = "amoghmahesh14@gmail.com"
    gmail_password = os.getenv("email_password")

    msg = MIMEMultipart()
    msg['From'] = gmail_user
    msg['To'] = to_email
    msg['Subject'] = subject

    # Attach the message
    msg.attach(MIMEText(message, 'plain'))

    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(gmail_user, gmail_password)
        text = msg.as_string()
        server.sendmail(gmail_user, to_email, text)
        server.quit()
        print("Email sent successfully!")
    except Exception as e:
        print("Email could not be sent. Error: {}".format(str(e)))

def email_reminders():
    users = Users.objects({})
    for user in users:
        current_applications =  user["applications"]
        for application in current_applications:
            if application["status"] != 3 or application["status"] != 4:
                try:
                    # Send an email reminder
                    to_email = user.email  # Use the user's email address
                    subject = "Job Application Reminder"
                    message = f"Hello {user.fullName},\n\nThe following job application has not been submitted yet.\nJob Title: {application['jobTitle']}\nCompany: {application['companyName']}\nApply By: {application['date']}\n\nBest regards,\nYour Application Tracker"
                    send_email(to_email, subject, message)
                except:
                    return jsonify({"error": "EMAIL wasn't sent"}), 400
                
sched = BackgroundScheduler(daemon=True)
sched.add_job(email_reminders,'interval',minutes=60)
sched.start()
def generate_pdf(data):
    doc = Document()

    # Set page margins to fit within one page
    sections = doc.sections
    for section in sections:
        section.left_margin = Pt(36)  # 0.5 inch
        section.right_margin = Pt(36)  # 0.5 inch
        section.top_margin = Pt(36)  # 0.5 inch
        section.bottom_margin = Pt(36)  # 0.5 inch

    # Helper function to add heading with format
    def add_heading_with_format(doc, text, font_size=16, is_bold=True):
        p = doc.add_paragraph()
        run = p.add_run(text)
        if is_bold:
            run.bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run.font.size = Pt(font_size)

    # Function to add details section
    def add_details_section(doc, section_title, details, is_bold_title=True):
        if section_title:
            add_heading_with_format(doc, section_title, font_size=14, is_bold=True)
        for detail in details:
            for key, value in detail.items():
                if key == "company":
                    p = doc.add_paragraph()
                    run = p.add_run(value)
                    run.bold = True
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                elif key == "project_title":
                    # Add the value of "project_title" with bold formatting
                    p = doc.add_paragraph()
                    run = p.add_run(value)
                    run.bold = True
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                elif key == "descriptionc":
                    # Add the value of "descriptionc" without "descriptionc" prefix
                    doc.add_paragraph(value, style="List Bullet")
                elif key != "descriptionc" and key != "level" and key != "extracurricularActivities":
                    if key == "university":
                        # Add the value of "university" with bold formatting and without a bullet
                        p = doc.add_paragraph()
                        run = p.add_run("University: " + value)
                        run.bold = True
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    else:
                        doc.add_paragraph(f"{value}", style="List Bullet")

    # Title
    add_heading_with_format(doc, "Resume", font_size=18, is_bold=True)

    # Contact Information
    add_heading_with_format(doc, "Contact Information", font_size=16, is_bold=True)
    doc.add_paragraph("Name: " + data["name"])
    doc.add_paragraph("Address: " + data["address"])
    doc.add_paragraph("Email: " + data["email"])
    doc.add_paragraph("LinkedIn: " + data["linkedin"])
    doc.add_paragraph("Phone: " + data["phone"])

    # Education section
    add_details_section(doc, "Education", data["education"])

    # Skills section
    skills = data["skills"]
    skills_text = ", ".join(skill["skills"] for skill in skills)
    add_heading_with_format(doc, "Skills", font_size=14, is_bold=True)
    doc.add_paragraph(skills_text, style="List Bullet")

    # Work Experience section
    add_heading_with_format(doc, "Work Experience", font_size=16, is_bold=True)
    for entry in data["workExperience"]:
        add_details_section(doc, "", [entry], is_bold_title=False)  # Removed the "Work Entry" heading

    # Projects section
    add_heading_with_format(doc, "Projects", font_size=16, is_bold=True)
    for project in data["projects"]:
        add_details_section(doc, "", [project], is_bold_title=False)  # Removed repeated "Project" heading

    # Save the document to a .docx file

    word_buffer = BytesIO()
    output_file_path = "generated_resume.docx"
    doc.save(word_buffer)
    word_buffer.seek(0)

    return word_buffer


@app.route('/resumebuilder', methods=['POST'])
def form_builder():
    try:
        # Assuming the request data is in JSON format
        data = request.json

        # Log the data (you can customize this part)
        print("Received Form Data:")
        for key, value in data.items():
            print(f"{key}: {value}")

        # Generate PDF
        pdf_data = generate_pdf(data)

        # Send the PDF file as a response
        return send_file(pdf_data, mimetype='application/msword', as_attachment=True,
                         attachment_filename='generated_resume.docx')
    except Exception as e:
        print(f"Error processing form data: {str(e)}")
        return "Error processing form data", 500


def get_job_app_status(applications):
    job_app_status = {
        "Jobs Saved": 0,
        "Applications": 0,
        "Interviews": 0,
        "Offers": 0,
    }
    for application in applications:
        # if application["status"]=="Job Saved":
        if int(application["status"]) == 1:
            job_app_status["Jobs Saved"] += 1
        # elif application["status"]=="Applications":
        elif int(application["status"]) == 2:
            job_app_status["Applications"] += 1
            job_app_status["Jobs Saved"] += 1
        # elif application["status"]=="Interviews":
        elif int(application["status"]) == 3:
            job_app_status["Interviews"] += 1
            job_app_status["Applications"] += 1
            job_app_status["Jobs Saved"] += 1
        # elif application["status"]=="Offers":
        elif int(application["status"]) == 4:
            job_app_status["Offers"] += 1
            job_app_status["Interviews"] += 1
            job_app_status["Applications"] += 1
            job_app_status["Jobs Saved"] += 1
    res = [
        {"name": "Jobs Saved", "count": job_app_status["Jobs Saved"]},
        {"name": "Applications", "count": job_app_status["Applications"]},
        {"name": "Interviews", "count": job_app_status["Interviews"]},
        {"name": "Offers", "count": job_app_status["Offers"]},
    ]
    return res,job_app_status["Applications"],job_app_status["Interviews"]

def get_last_six_months_job_counts(applications):
    month_map = {
        index + 1: val
        for index, val in enumerate(
            [
                "Jan", "Feb", "Mar", "Apr",
                "May", "Jun", "Jul", "Aug",
                "Sep", "Oct", "Nov", "Dec",
            ]
        )
    }
    # Create a defaultdict to store date strings for each month
    result_dict = defaultdict(list)
    # Get the current date
    current_date = (datetime.now().replace(day=1) + timedelta(days=32)).replace(day=1)
    # Iterate over the last six months
    for i in range(6):
        # Calculate the start and end date of the current month
        end_of_month = current_date.replace(day=1) - timedelta(days=1)
        start_of_month = end_of_month.replace(day=1)
        # Filter date objects that belong to the current month
        current_month_dates = [
            datetime.strptime(application["date"], "%Y-%m-%d").strftime("%Y-%m-%d")
            for application in applications
            if start_of_month
            <= datetime.strptime(application["date"], "%Y-%m-%d")
            <= end_of_month
        ]
        # Store the result in the dictionary
        result_dict[
            "%s %s" % (month_map[start_of_month.month], start_of_month.year)
        ] = len(current_month_dates)
        # Move to the previous month
        current_date = start_of_month
    res = [
        {"Month": key, "Jobs Created": result_dict[key]}
        for key in list(result_dict.keys())[::-1]
    ]
    return res

def get_last_four_jobs(applications):
    # apps = defaultdict(list)
    # for application in applications:
    #     if datetime.strptime(application["date"], "%Y-%m-%d") not in apps:
    #         apps[datetime.strptime(application["date"], "%Y-%m-%d")] = []
    #     apps[datetime.strptime(application["date"], "%Y-%m-%d")].append(application)
    appStatus = {'1': "Job Saved", '2': "Applied", '3': "Interviewed", '4':"Offered"}
    apps = sorted(applications, key=lambda application: datetime.strptime(application["date"], "%Y-%m-%d"), reverse=True)[:4]
    for app in apps:
        print(app)
    res = [
        {
            "jobTitle": application["jobTitle"],
            "company": application["companyName"],
            "status": appStatus[application["status"]],
        }
        for application in apps
    ]
    print(res)
    return res

if __name__ == '__main__':
    app.run(debug=False)

# if __name__ == "__main__":
#     app.run()